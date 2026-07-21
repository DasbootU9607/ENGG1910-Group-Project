from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parent
REPORT_MD = PROJECT_DIR / "report_draft.md"
REPORT_DOCX = PROJECT_DIR / "report_draft.docx"
CHART_PATH = PROJECT_DIR / "outputs" / "risk_distribution.png"


def set_font(run, name: str = "Calibri", size: int = 11, bold: bool = False, color: str = "000000") -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    set_font(run)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.space_before = Pt(12 if level == 2 else 16)
    paragraph.paragraph_format.space_after = Pt(6 if level == 2 else 8)
    run = paragraph.add_run(text)
    set_font(run, size=13 if level == 2 else 16, bold=True, color="2E74B5")


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(text)
    set_font(run, size=16, bold=True, color="0B2545")


def build_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    lines = REPORT_MD.read_text(encoding="utf-8").splitlines()
    skip_next_image = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            add_title(document, stripped[2:])
        elif stripped.startswith("## Project Title:"):
            add_title(document, stripped[3:])
        elif stripped.startswith("## "):
            add_heading(document, stripped[3:], 1)
        elif stripped.startswith("!["):
            if CHART_PATH.exists() and not skip_next_image:
                document.add_picture(str(CHART_PATH), width=Inches(5.9))
                last = document.paragraphs[-1]
                last.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption = document.add_paragraph("Figure 1. Risk level distribution on the LIAR test set.")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(8)
                for run in caption.runs:
                    set_font(run, size=10, color="555555")
            skip_next_image = True
        elif stripped.startswith("[") and "] " in stripped:
            add_paragraph(document, stripped)
        else:
            add_paragraph(document, stripped.replace("`", ""))

    REPORT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(REPORT_DOCX)
    print(REPORT_DOCX)


if __name__ == "__main__":
    build_docx()
